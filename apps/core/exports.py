from __future__ import annotations

import mimetypes
import re
from datetime import datetime
from decimal import Decimal
from pathlib import Path
from typing import Any, Iterable
from zipfile import ZIP_DEFLATED, ZipFile
from xml.sax.saxutils import escape as xml_escape
from copy import copy

from openpyxl import load_workbook

from django.conf import settings

from .models import (
    DocumentRecord,
    PPEIssuance,
    PrimaryDocument,
    ProcurementRequest,
    SiteMaterialRequest,
    SMRContract,
    StockIssue,
    StockReceipt,
    SupplierDocument,
    SupplyContract,
    WorkAcceptanceAct,
    WriteOffAct,
    WriteOffTemplateVariant,
)
from .reporting import REPORT_PROVIDERS


def money(value: Any) -> str:
    numeric = float(value or 0)
    return f"{numeric:,.2f}".replace(",", " ").replace(".", ",")


PLACEHOLDER_RE = re.compile(r"\{\{\s*([A-Z0-9_]+)\s*\}\}")
CELL_REF_RE = re.compile(r'([A-Z]+)(\d+)')
MONTH_NAMES = [
    "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]

DOCX_TEMPLATE_FILES = {
    "smr_contract": "Договор на СМР_шаблон.docx",
    "supply_contract": "Договор поставки_шаблон.docx",
    "stock_receipt": "Приходный ордер_шаблон.docx",
    "stock_issue": "Требование-накладная_шаблон.docx",
    "write_off": "Акт списания материалов по договору_шаблон.docx",
    "write_off_production_economic": "Акт списания материалов на производственно-хозяйственные нужды_шаблон.docx",
    "work_acceptance": "Акт сдачи-приемки выполненных работ_шаблон.docx",
}

PRIMARY_DOCUMENT_TEMPLATE_FILES = {
    "invoice": "Счет на материал_шаблон.docx",
    "invoice_facture": "Счет-фактура_шаблон.docx",
    "vat_invoice": "Счет-фактура_шаблон.docx",
    "goods_waybill": "Товарная накладная ТОРГ-12_шаблон.docx",
    "upd": "Товарная накладная ТОРГ-12_шаблон.docx",
    "receipt_invoice": "Товарная накладная ТОРГ-12_шаблон.docx",
    "payment_order": "Платежное поручение_шаблон.docx",
}

PRIMARY_DOCUMENT_XLSX_TEMPLATE_FILES = {
    "invoice": "Счет на оплату_шаблон.xlsx",
}

XLSX_TEMPLATE_FILES = {
    "site_material_report": "Материальный отчет_шаблон.xlsx",
}


def _load_xlsxwriter():
    try:
        import xlsxwriter
    except ImportError as exc:
        raise RuntimeError("Экспорт XLSX недоступен: пакет xlsxwriter не установлен или поврежден.") from exc
    return xlsxwriter


def _load_docx_dependencies():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError("Экспорт DOCX недоступен: пакет python-docx или lxml не установлен либо поврежден.") from exc
    return Document, WD_ALIGN_PARAGRAPH, Pt


class Exporter:

    def _short_name(self, full_name: str) -> str:
        """Возвращает формат И.О.Фамилия"""
        parts = (full_name or "").strip().split()
        if len(parts) >= 3:
            return f"{parts[1][0]}.{parts[2][0]}.{parts[0]}"
        elif len(parts) == 2:
            return f"{parts[1][0]}.{parts[0]}"
        return full_name or "________________"

    def _last_name_initials(self, full_name: str) -> str:
        """Возвращает формат Фамилия И.О. (для комиссии в актах)"""
        parts = (full_name or "").strip().split()
        if len(parts) >= 3:
            return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
        elif len(parts) == 2:
            return f"{parts[0]} {parts[1][0]}."
        return full_name or "________________"

    def _warehouse_user_name(self) -> str:
        from .models import User, RoleChoices
        user = User.objects.filter(role=RoleChoices.WAREHOUSE, is_active=True).first()
        return user.full_name_or_username if user else "________________"

    def _export_work_schedule(self, entity_id: int) -> Path:
        from .models import WorkSchedule
        Document, WD_ALIGN_PARAGRAPH, Pt = _load_docx_dependencies()
        from docx.shared import Cm
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        schedule = WorkSchedule.objects.select_related("contract", "created_by").prefetch_related("lines").get(pk=entity_id)

        profile = self._organization_profile()
        org_name = profile["name"] or "АО «СТ-1»"
        creator_name = schedule.created_by.full_name_or_username if schedule.created_by_id else ""
        creator_short = self._short_name(creator_name) if creator_name else "________________"
        year = schedule.period_start.year if schedule.period_start else ""
        object_name = schedule.contract.object.name if schedule.contract.object else "-"
        customer_name = schedule.contract.resolved_customer_name() or "-"

        customer_signer = schedule.contract.customer_signer_name or ""
        customer_signer_short = self._short_name(customer_signer) if customer_signer else "________________"

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(10)

        header_table = doc.add_table(rows=1, cols=2)
        left = header_table.cell(0, 0)
        right = header_table.cell(0, 1)
        self._hide_table_borders(header_table)

        left.text = ""
        left_para = left.paragraphs[0]
        left_para.add_run("СОГЛАСОВАНО\n").bold = True
        left_para.add_run(f"Заместитель генерального директора\n{org_name}\n\n_____________________ {creator_short}")

        right.text = ""
        right_para = right.paragraphs[0]
        right_para.add_run("УТВЕРЖДАЮ\n").bold = True
        right_para.add_run(f"Директор программы\n{customer_name}\n\n_____________________ {customer_signer_short}")

        doc.add_paragraph()

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("ГРАФИК")
        title_run.bold = True
        title_run.font.size = Pt(12)

        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_para.add_run(
            f"производства работ с {self._date_text(schedule.period_start)} по {self._date_text(schedule.period_end)}\n"
            f"по договору СМР № {schedule.contract.number}, объект: {object_name}"
        )

        doc.add_paragraph()

        lines = list(schedule.lines.all().order_by("order"))
        self._add_table(
            doc,
            ["№ п/п", "Наименование работ", "Этап", "Исполнитель", "Дата начала", "Дата окончания", "Примечание"],
            [
                [
                    str(line.order),
                    line.work_type,
                    line.stage,
                    line.executor or creator_name,
                    self._date_text(line.start_date),
                    self._date_text(line.end_date),
                    line.notes,
                ]
                for line in lines
            ],
        )

        doc.add_paragraph()

        sign_table = doc.add_table(rows=1, cols=2)
        sign_left = sign_table.cell(0, 0)
        sign_right = sign_table.cell(0, 1)
        self._hide_table_borders(sign_table)

        sign_left.text = ""
        sign_left.paragraphs[0].add_run(
            f"От {org_name}\nНачальник участка {schedule.site_name}\n\n_____________________ {creator_short}"
        )

        sign_right.text = ""
        sign_right.paragraphs[0].add_run(
            f"От {customer_name}\n\n\n_____________________ {customer_signer_short}"
        )

        path = self._doc_path("work_schedule", schedule.number)
        doc.save(path)
        return path

    def content_type(self, path: Path) -> str:
        guessed, _ = mimetypes.guess_type(path.name)
        return guessed or "application/octet-stream"

    def _template_path(self, template_name: str) -> Path | None:
        templates_dir = Path(getattr(settings, "DOCUMENT_TEMPLATES_DIR", ""))
        path = templates_dir / template_name
        return path if path.exists() else None

    def _render_docx_template(self, template_name: str, context: dict[str, Any], path: Path) -> bool:
        template_path = self._template_path(template_name)
        if not template_path:
            return False
        replacements = {key: "" if value is None else str(value) for key, value in context.items()}
        def replace_placeholders(raw_xml: bytes) -> bytes:
            xml_text = raw_xml.decode("utf-8")
            def replace(match: re.Match[str]) -> str:
                return xml_escape(replacements.get(match.group(1), ""))
            return PLACEHOLDER_RE.sub(replace, xml_text).encode("utf-8")
        with ZipFile(template_path, "r") as source, ZipFile(path, "w", ZIP_DEFLATED) as target:
            for member in source.infolist():
                data = source.read(member.filename)
                if member.filename.startswith("word/") and member.filename.endswith(".xml"):
                    data = replace_placeholders(data)
                target.writestr(member, data)
        return True

    def _render_docx_template_with_table_rows(
        self,
        template_name: str,
        context: dict[str, Any],
        path: Path,
        *,
        table_rows: list[dict[str, Any]] | None = None,
        row_marker: str = "",
    ) -> bool:
        template_path = self._template_path(template_name)
        if not template_path:
            return False
        replacements = {key: "" if value is None else str(value) for key, value in context.items()}
        def expand_table_rows(xml_text: str) -> str:
            if not table_rows or not row_marker:
                return xml_text
            marker_pos = xml_text.find(row_marker)
            if marker_pos == -1:
                return xml_text
            tr_start = xml_text.rfind("<w:tr ", 0, marker_pos)
            if tr_start == -1:
                tr_start = xml_text.rfind("<w:tr>", 0, marker_pos)
            if tr_start == -1:
                return xml_text
            tr_end_marker = "</w:tr>"
            tr_end = xml_text.find(tr_end_marker, marker_pos)
            if tr_end == -1:
                return xml_text
            tr_end += len(tr_end_marker)
            row_template = xml_text[tr_start:tr_end]
            expanded_rows = []
            for row_data in table_rows:
                row_xml = row_template
                for placeholder_key, placeholder_value in row_data.items():
                    marker = "{{" + placeholder_key + "}}"
                    row_xml = row_xml.replace(marker, xml_escape(str(placeholder_value or "")))
                expanded_rows.append(row_xml)
            return xml_text[:tr_start] + "".join(expanded_rows) + xml_text[tr_end:]
        def replace_placeholders(raw_xml: bytes) -> bytes:
            xml_text = raw_xml.decode("utf-8")
            xml_text = expand_table_rows(xml_text)
            def replace(match: re.Match[str]) -> str:
                return xml_escape(replacements.get(match.group(1), ""))
            return PLACEHOLDER_RE.sub(replace, xml_text).encode("utf-8")
        with ZipFile(template_path, "r") as source, ZipFile(path, "w", ZIP_DEFLATED) as target:
            for member in source.infolist():
                data = source.read(member.filename)
                if member.filename.startswith("word/") and member.filename.endswith(".xml"):
                    data = replace_placeholders(data)
                target.writestr(member, data)
        return True

    def _render_docx_file_with_table_rows(
        self,
        source_path: Path,
        context: dict[str, Any],
        path: Path,
        *,
        table_rows: list[dict[str, Any]] | None = None,
        row_marker: str = "",
    ) -> bool:
        if not source_path.exists():
            return False
        replacements = {key: "" if value is None else str(value) for key, value in context.items()}
        def expand_table_rows(xml_text: str) -> str:
            if not table_rows or not row_marker:
                return xml_text
            marker_pos = xml_text.find(row_marker)
            if marker_pos == -1:
                return xml_text
            tr_start = xml_text.rfind("<w:tr ", 0, marker_pos)
            if tr_start == -1:
                tr_start = xml_text.rfind("<w:tr>", 0, marker_pos)
            if tr_start == -1:
                return xml_text
            tr_end_marker = "</w:tr>"
            tr_end = xml_text.find(tr_end_marker, marker_pos)
            if tr_end == -1:
                return xml_text
            tr_end += len(tr_end_marker)
            row_template = xml_text[tr_start:tr_end]
            expanded_rows = []
            for row_data in table_rows:
                row_xml = row_template
                for placeholder_key, placeholder_value in row_data.items():
                    marker = "{{" + placeholder_key + "}}"
                    row_xml = row_xml.replace(marker, xml_escape(str(placeholder_value or "")))
                expanded_rows.append(row_xml)
            return xml_text[:tr_start] + "".join(expanded_rows) + xml_text[tr_end:]
        def replace_placeholders(raw_xml: bytes) -> bytes:
            xml_text = raw_xml.decode("utf-8")
            xml_text = expand_table_rows(xml_text)
            def replace(match: re.Match[str]) -> str:
                return xml_escape(replacements.get(match.group(1), ""))
            return PLACEHOLDER_RE.sub(replace, xml_text).encode("utf-8")
        with ZipFile(source_path, "r") as source, ZipFile(path, "w", ZIP_DEFLATED) as target:
            for member in source.infolist():
                data = source.read(member.filename)
                if member.filename.startswith("word/") and member.filename.endswith(".xml"):
                    data = replace_placeholders(data)
                target.writestr(member, data)
        return True

    def _render_xlsx_template(self, template_name: str, context: dict[str, Any], path: Path) -> bool:
        template_path = self._template_path(template_name)
        if not template_path:
            return False
        replacements = {key: "" if value is None else str(value) for key, value in context.items()}
        def replace_placeholders(raw_xml: bytes) -> bytes:
            xml_text = raw_xml.decode("utf-8")
            def replace(match: re.Match[str]) -> str:
                return xml_escape(replacements.get(match.group(1), ""))
            return PLACEHOLDER_RE.sub(replace, xml_text).encode("utf-8")
        with ZipFile(template_path, "r") as source, ZipFile(path, "w", ZIP_DEFLATED) as target:
            for member in source.infolist():
                data = source.read(member.filename)
                if member.filename.endswith(".xml"):
                    data = replace_placeholders(data)
                target.writestr(member, data)
        return True

    def _render_xlsx_template_with_rows(
        self,
        template_name: str,
        context: dict[str, Any],
        path: Path,
        *,
        table_rows: list[dict[str, Any]],
        template_row_number: int,
    ) -> bool:
        template_path = self._template_path(template_name)
        if not template_path:
            return False

        wb = load_workbook(template_path)
        ws = wb.active
        max_col = ws.max_column

        rows = table_rows or []
        n = len(rows)
        delta = n - 1

        template_styles = []
        template_values = []
        for col in range(1, max_col + 1):
            cell = ws.cell(row=template_row_number, column=col)
            template_styles.append({
                "font": copy(cell.font),
                "border": copy(cell.border),
                "fill": copy(cell.fill),
                "alignment": copy(cell.alignment),
                "number_format": cell.number_format,
            })
            template_values.append(cell.value)

        merges_below = []
        for merged_range in list(ws.merged_cells.ranges):
            if merged_range.min_row > template_row_number:
                merges_below.append(
                    (merged_range.min_col, merged_range.min_row, merged_range.max_col, merged_range.max_row)
                )
                ws.unmerge_cells(str(merged_range))

        if delta > 0:
            ws.insert_rows(template_row_number + 1, amount=delta)
        elif delta < 0:
            ws.delete_rows(template_row_number, amount=1)

        for i, row_data in enumerate(rows):
            row_idx = template_row_number + i
            for col in range(1, max_col + 1):
                cell = ws.cell(row=row_idx, column=col)
                style = template_styles[col - 1]
                cell.font = style["font"]
                cell.border = style["border"]
                cell.fill = style["fill"]
                cell.alignment = style["alignment"]
                cell.number_format = style["number_format"]

                value = template_values[col - 1]
                if isinstance(value, str):
                    for key, val in row_data.items():
                        marker = "{{" + key + "}}"
                        if marker in value:
                            value = value.replace(marker, "" if val is None else str(val))
                    value = PLACEHOLDER_RE.sub("", value)
                    cell.value = value if value != "" else None
                else:
                    cell.value = value

        for min_col, min_row, max_col_m, max_row in merges_below:
            new_min_row = min_row + delta
            new_max_row = max_row + delta
            ws.merge_cells(
                start_row=new_min_row, start_column=min_col,
                end_row=new_max_row, end_column=max_col_m,
            )

        replacements = {k: "" if v is None else str(v) for k, v in context.items()}
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and "{{" in cell.value:
                    cell.value = PLACEHOLDER_RE.sub(lambda m: replacements.get(m.group(1), ""), cell.value)

        wb.save(path)
        return True

    def _date_text(self, value: Any) -> str:
        if not value:
            return ""
        return value.strftime("%d.%m.%Y") if hasattr(value, "strftime") else str(value)

    def _date_parts(self, prefix: str, value: Any) -> dict[str, str]:
        if not value or not hasattr(value, "month"):
            return {f"{prefix}_DAY": "", f"{prefix}_MONTH": "", f"{prefix}_YEAR": ""}
        return {
            f"{prefix}_DAY": f"{value.day:02d}",
            f"{prefix}_MONTH": MONTH_NAMES[value.month - 1],
            f"{prefix}_YEAR": str(value.year),
        }

    def _duration_days(self, start_date: Any, end_date: Any) -> str:
        if not start_date or not end_date:
            return ""
        try:
            return str((end_date - start_date).days + 1)
        except TypeError:
            return ""

    def _template_common_context(self) -> dict[str, str]:
        profile = self._organization_profile()
        organization_name = profile["name"] or "АО «СТ-1»"
        return {
            "ORGANIZATION_NAME": organization_name,
            "CITY": "Москва",
            "CONTRACT_CITY": "Москва",
            "DIRECTOR_NAME": "________________",
            "RESPONSIBLE_PERSON_NAME": "________________",
            "SITE_MANAGER_NAME": "________________",
            "SITE_MANAGER_NAME_SHORT": "________________",
            "RESPONSIBLE_PERSON_NAME_SHORT": "________________",
            "LEFT_SIGNER_NAME": "________________",
            "RIGHT_SIGNER_NAME": "________________",
            "CONTRACTOR_SIGNER_NAME": "________________",
            "CONTRACTOR_SIGNER_POSITION": "представителя",
            "CONTRACTOR_AUTH_DOC": "доверенности",
            "CUSTOMER_SIGNER_NAME": "________________",
            "CUSTOMER_SIGNER_POSITION": "представителя",
            "CUSTOMER_AUTH_DOC": "доверенности",
            "CUSTOMER_SIGNER_AUTH_DOC": "доверенности",
            "BUYER_SIGNER": "________________",
            "BUYER_AUTH_DOC": "доверенности",
            "SUPPLIER_SIGNER": "________________",
            "SUPPLIER_AUTH_DOC": "доверенности",
            "JURISDICTION_PARTY": organization_name,
        }

    def _add_line_context(self, context: dict[str, Any], lines: Iterable[Any], mapper, *, limit: int = 12) -> None:
        for index, line in enumerate(list(lines)[:limit], start=1):
            suffix = "" if index == 1 else f"_{index}"
            for key, value in mapper(line, index).items():
                context[f"{key}{suffix}"] = value

    def _line_amount(self, line: Any) -> Decimal:
        return Decimal(line.quantity or 0) * Decimal(line.unit_price or 0)

    def _xlsx_path(self, prefix: str, date_from: Any, date_to: Any) -> Path:
        return settings.EXPORTS_DIR / f"{prefix}_{date_from}_{date_to}.xlsx"

    def _site_material_report_template_context(self, rows: list[dict[str, Any]], filters: dict[str, Any], *, user=None) -> dict[str, Any]:
        date_from = filters.get("date_from") or datetime.now().date().replace(day=1)
        date_to = filters.get("date_to") or datetime.now().date()
        detail_rows = [row for row in rows if str(row.get("Период", "")).upper() != "ИТОГО"]
        site_name = (
            (getattr(user, "site_name", "") or "").strip()
            or (filters.get("location_name") or "").strip()
            or (detail_rows[0].get("Место хранения", "") if detail_rows else "")
        )
        total_opening = sum(float(row.get("Остаток на начало", 0) or 0) for row in detail_rows)
        total_receipt = sum(float(row.get("Поступило за период", 0) or 0) for row in detail_rows)
        total_issue = sum(float(row.get("Израсходовано за период", 0) or 0) for row in detail_rows)
        total_closing = sum(float(row.get("Остаток на конец", 0) or 0) for row in detail_rows)
        total_amount = sum(float(row.get("Сумма остатка", 0) or 0) for row in detail_rows)
        context: dict[str, Any] = {
            **self._template_common_context(),
            "REPORT_NUMBER": f"MAT-{date_from:%Y%m}",
            "REPORT_PERIOD": f"{self._date_text(date_from)} - {self._date_text(date_to)}",
            "SITE_NAME": site_name or "-",
            "RESPONSIBLE_PERSON_NAME": getattr(user, "full_name_or_username", "") if user is not None else "",
            "TOTAL_OPENING": total_opening,
            "TOTAL_RECEIPT": total_receipt,
            "TOTAL_ISSUE": total_issue,
            "TOTAL_CLOSING": total_closing,
            "TOTAL_AMOUNT": total_amount,
        }
        lines_data = []
        for index, row in enumerate(detail_rows, start=1):
            lines_data.append({
                "LINE_NO": str(index),
                "MATERIAL_CODE": row.get("Код материала", ""),
                "MATERIAL_NAME": row.get("Наименование материала", ""),
                "UNIT": row.get("Ед. изм.", ""),
                "OPENING_QTY": row.get("Остаток на начало", ""),
                "RECEIPT_QTY": row.get("Поступило за период", ""),
                "ISSUE_QTY": row.get("Израсходовано за период", ""),
                "CLOSING_QTY": row.get("Остаток на конец", ""),
                "PRICE": row.get("Цена за единицу", ""),
                "CLOSING_AMOUNT": row.get("Сумма остатка", ""),
                "BASIS_DOCUMENT": row.get("Период", ""),
            })
        context["__lines_data__"] = lines_data
        return context

    def export_document(self, entity_type: str, entity_id: int) -> Path:
        handlers = {
            "smr_contract": self._export_smr_contract,
            "supply_contract": self._export_supply_contract,
            "site_material_request": self._export_site_material_request,
            "procurement_request": self._export_procurement_request,
            "primary_document": self._export_primary_document,
            "stock_receipt": self._export_stock_receipt,
            "stock_issue": self._export_stock_issue,
            "write_off": self._export_writeoff,
            "ppe_issuance": self._export_ppe_issuance,
            "supplier_document": self._export_supplier_document,
            "work_acceptance": self._export_work_acceptance,
            "work_schedule": self._export_work_schedule,
        }
        if entity_type not in handlers:
            raise ValueError("Для этого типа документа выгрузка не реализована.")
        path = handlers[entity_type](entity_id)
        DocumentRecord.objects.filter(entity_type=entity_type, entity_id=entity_id).update(file_path=str(path))
        return path

    def export_report(self, report_name: str, filters: dict[str, Any], *, user=None) -> Path:
        provider = REPORT_PROVIDERS[report_name]
        rows = provider(filters, user=user)
        date_from = filters.get("date_from") or datetime.now().date().replace(day=1)
        date_to = filters.get("date_to") or datetime.now().date()
        path = self._xlsx_path(report_name, date_from, date_to)
        template_name = XLSX_TEMPLATE_FILES.get(report_name)
        if template_name:
            context = self._site_material_report_template_context(rows, filters, user=user)
            lines_data = context.pop("__lines_data__", [])
            if self._render_xlsx_template_with_rows(
                template_name, context, path,
                table_rows=lines_data, template_row_number=7,
            ):
                return path
        xlsxwriter = _load_xlsxwriter()
        workbook = xlsxwriter.Workbook(str(path))
        worksheet = workbook.add_worksheet("Отчет")
        title_format = workbook.add_format({"bold": True, "font_size": 14})
        header_format = workbook.add_format({"bold": True, "bg_color": "#E6D7C6", "border": 1})
        cell_format = workbook.add_format({"border": 1})
        numeric_format = workbook.add_format({"border": 1, "num_format": "#,##0.00"})
        total_cell_format = workbook.add_format({"border": 1, "bold": True, "bg_color": "#F3E7D7"})
        total_numeric_format = workbook.add_format({"border": 1, "bold": True, "bg_color": "#F3E7D7", "num_format": "#,##0.00"})
        worksheet.write("A1", "Экспорт отчета АИС", title_format)
        worksheet.write("A2", f"Период: {date_from} - {date_to}")
        worksheet.freeze_panes(4, 0)
        if rows:
            headers = list(rows[0].keys())
            for col_index, header in enumerate(headers):
                worksheet.write(3, col_index, header, header_format)
            for row_index, row in enumerate(rows, start=4):
                first_header = headers[0]
                first_value = row.get(first_header, "")
                is_total_row = str(first_value).startswith("ИТОГО")
                for col_index, header in enumerate(headers):
                    value = row.get(header)
                    if isinstance(value, Decimal):
                        value = float(value)
                    if isinstance(value, (int, float)):
                        worksheet.write_number(
                            row_index, col_index, float(value),
                            total_numeric_format if is_total_row else numeric_format,
                        )
                    else:
                        worksheet.write(
                            row_index, col_index, value,
                            total_cell_format if is_total_row else cell_format,
                        )
            for col_index, header in enumerate(headers):
                max_len = max(len(str(header)), *(len(str(item.get(header, ""))) for item in rows))
                worksheet.set_column(col_index, col_index, min(max_len + 2, 45))
            worksheet.autofilter(3, 0, 3 + len(rows), len(headers) - 1)
        else:
            worksheet.write("A4", "Нет данных за выбранный период.")
        workbook.close()
        return path

    def _export_path(self, prefix: str, number: str, extension: str) -> Path:
        safe_number = number.replace("/", "_").replace("\\", "_").replace(" ", "_")
        return settings.EXPORTS_DIR / f"{prefix}_{safe_number}.{extension}"

    def _doc_path(self, prefix: str, number: str) -> Path:
        return self._export_path(prefix, number, "docx")

    def _xlsx_document_path(self, prefix: str, number: str) -> Path:
        return self._export_path(prefix, number, "xlsx")

    def _prepare_doc(self, title: str, subtitle: str = ""):
        Document, WD_ALIGN_PARAGRAPH, Pt = _load_docx_dependencies()
        document = Document()
        style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(14)
        if subtitle:
            subtitle_paragraph = document.add_paragraph()
            subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_paragraph.add_run(subtitle).italic = True
        document.add_paragraph()
        return document

    def _add_meta(self, document, items: Iterable[tuple[str, str]]) -> None:
        for label, value in items:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{label}: ").bold = True
            paragraph.add_run(value)

    def _add_heading(self, document, text: str) -> None:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = _load_docx_dependencies()[2](12)

    def _add_clause(self, document, number: str, text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{number}. ").bold = True
        paragraph.add_run(text)

    def _add_table(self, document, headers: list[str], rows: list[list[str]]) -> None:
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = value

    def _add_signature(self, document, left_label: str, right_label: str) -> None:
        document.add_paragraph()
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.cell(0, 0).text = f"{left_label}\n\n_____________________"
        table.cell(0, 1).text = f"{right_label}\n\n_____________________"

    def _add_signature_no_border(self, document, left_label: str, right_label: str) -> None:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        document.add_paragraph()
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = OxmlElement("w:tblBorders")
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tblBorders.append(border)
        tblPr.append(tblBorders)
        table.cell(0, 0).text = left_label
        table.cell(0, 1).text = right_label
    def _hide_table_borders(self, table) -> None:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = OxmlElement("w:tblBorders")
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tblBorders.append(border)
        tblPr.append(tblBorders)
    def _add_signature_borderless(self, document, left_label: str, right_label: str, left_name: str, right_name: str) -> None:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        document.add_paragraph()
        table = document.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = OxmlElement("w:tblBorders")
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tblBorders.append(border)
        tblPr.append(tblBorders)
        table.cell(0, 0).text = left_label
        table.cell(0, 1).text = right_label
        table.cell(1, 0).text = f"_____________________ / {left_name} /"
        table.cell(1, 1).text = f"_____________________ / {right_name} /"

    def _organization_profile(self) -> dict[str, str]:
        from .models import OrganizationProfile
        profile = OrganizationProfile.get()
        return {
            "name": profile.name,
            "tax_id": profile.tax_id,
            "kpp": profile.kpp,
            "ogrn": profile.ogrn,
            "address": profile.address,
            "bank_details": profile.bank_details,
            "requisites": profile.requisites,
            "contractor_signer_name": profile.contractor_signer_name,
            "contractor_signer_position": profile.contractor_signer_position,
            "contractor_signer_name_genitive": profile.contractor_signer_name_genitive,
            "contractor_signer_position_genitive": profile.contractor_signer_position_genitive,
            "contractor_auth_doc": profile.contractor_auth_doc,
            "bank_name": profile.bank_name,
            "bik": profile.bik,
            "account": profile.account,
            "corr_account": profile.corr_account,
            "okpo": profile.okpo,
        }

    def _organization_name(self) -> str:
        return self._organization_profile()["name"]

    def _organization_requisites(self) -> str:
        profile = self._organization_profile()
        if profile["requisites"]:
            return profile["requisites"]
        parts: list[str] = []
        if profile["tax_id"]:
            parts.append(f"ИНН {profile['tax_id']}")
        if profile["kpp"]:
            parts.append(f"КПП {profile['kpp']}")
        if profile["ogrn"]:
            parts.append(f"ОГРН {profile['ogrn']}")
        if profile["address"]:
            parts.append(profile["address"])
        if profile["bank_details"]:
            parts.append(profile["bank_details"])
        return "; ".join(part for part in parts if part)

    def _supplier_requisites(self, supplier) -> str:
        if hasattr(supplier, "requisites_text"):
            return supplier.requisites_text()
        return ""

    def _extract_requisite(self, text: str, pattern: str) -> str:
        if not text:
            return ""
        match = re.search(pattern, text, flags=re.IGNORECASE)
        return match.group(1).strip() if match else ""

    def _smr_contract_template_context(self, contract: SMRContract) -> dict[str, Any]:
        object_name = contract.object.name if contract.object else ""
        vat_amount = Decimal(contract.amount or 0) * Decimal(contract.vat_rate or 0) / Decimal("100")
        profile = self._organization_profile()
        customer_signer = contract.customer_signer_name or ""
        contractor_signer = profile.get("contractor_signer_name") or ""
        context = {
            **self._template_common_context(),
            **self._date_parts("CONTRACT", contract.contract_date),
            "CONTRACT_NUMBER": contract.number,
            "CUSTOMER_NAME": contract.resolved_customer_name() or "-",
            "CUSTOMER_REQUISITES": contract.resolved_customer_requisites() or "-",
            "CUSTOMER_NAME_SHORT": (contract.object.customer_name_short if contract.object else "") or contract.resolved_customer_name() or "-",
            "CUSTOMER_LEGAL_ADDRESS": (contract.object.customer_legal_address if contract.object else "") or "",
            "CUSTOMER_TAX_ID": (contract.object.customer_tax_id if contract.object else "") or "",
            "CUSTOMER_KPP": (contract.object.customer_kpp if contract.object else "") or "",
            "CUSTOMER_OGRN": (contract.object.customer_ogrn if contract.object else "") or "",
            "CUSTOMER_BANK": (contract.object.customer_bank if contract.object else "") or "",
            "CUSTOMER_BIK": (contract.object.customer_bik if contract.object else "") or "",
            "CUSTOMER_ACCOUNT": (contract.object.customer_account if contract.object else "") or "",
            "CUSTOMER_CORR_ACCOUNT": (contract.object.customer_corr_account if contract.object else "") or "",
            "CUSTOMER_OKPO": (contract.object.customer_okpo if contract.object else "") or "",
            "CUSTOMER_SIGNER_NAME_GENITIVE": contract.customer_signer_name_genitive or contract.customer_signer_name or "________________",
            "CUSTOMER_SIGNER_POSITION_GENITIVE": contract.customer_signer_position_genitive or contract.customer_signer_position or "представителя",
            "CONTRACTOR_SIGNER_NAME_GENITIVE": profile.get("contractor_signer_name_genitive") or contractor_signer or "________________",
            "CONTRACTOR_SIGNER_POSITION_GENITIVE": profile.get("contractor_signer_position_genitive") or profile.get("contractor_signer_position") or "представителя",
            "CONTRACTOR_NAME": contract.resolved_contractor_name() or profile.get("name") or "-",
            "CONTRACTOR_REQUISITES": contract.resolved_contractor_requisites() or "-",
            "CONTRACTOR_ADDRESS": profile.get("address") or "",
            "CONTRACTOR_TAX_ID": profile.get("tax_id") or "",
            "CONTRACTOR_KPP": profile.get("kpp") or "",
            "CONTRACTOR_OGRN": profile.get("ogrn") or "",
            "CONTRACTOR_BANK": profile.get("bank_name") or "",
            "CONTRACTOR_BIK": profile.get("bik") or "",
            "CONTRACTOR_ACCOUNT": profile.get("account") or "",
            "CONTRACTOR_CORR_ACCOUNT": profile.get("corr_account") or "",
            "CONTRACTOR_OKPO": profile.get("okpo") or "",
            "OBJECT_NAME": object_name,
            "OBJECT_ADDRESS": contract.object.address if contract.object else "",
            "WORK_SUBJECT": contract.subject,
            "WORK_LINE_NO": "1",
            "WORK_NAME": contract.work_type or contract.subject,
            "WORK_QUANTITY": contract.planned_volume or "",
            "WORK_UNIT": contract.volume_unit or "",
            "ESTIMATE_LINE": f"{contract.work_type or contract.subject}: {contract.planned_volume or '-'} {contract.volume_unit or ''}".strip(),
            "CONTRACT_AMOUNT_TEXT": f"{money(contract.amount)} руб.",
            "VAT_AMOUNT_TEXT": f"{money(vat_amount)} руб.",
            "WORK_DURATION_DAYS": self._duration_days(contract.start_date, contract.end_date),
            "CUSTOMER_SIGNER_NAME": customer_signer or "________________",
            "CUSTOMER_SIGNER_NAME_SHORT": self._short_name(customer_signer) if customer_signer else "________________",
            "CUSTOMER_SIGNER_NAME_FULL": customer_signer or "________________",
            "CUSTOMER_SIGNER_POSITION": contract.customer_signer_position or "представителя",
            "CUSTOMER_AUTH_DOC": contract.customer_auth_doc or "доверенности",
            "CONTRACTOR_SIGNER_NAME": contractor_signer or "________________",
            "CONTRACTOR_SIGNER_NAME_SHORT": self._short_name(contractor_signer) if contractor_signer else "________________",
            "CONTRACTOR_SIGNER_POSITION": profile.get("contractor_signer_position") or "представителя",
            "CONTRACTOR_AUTH_DOC": profile.get("contractor_auth_doc") or "доверенности",
            "WORK_OBJECT_DESCRIPTION": contract.work_object_description or "",
            "WORK_BASIS_TEXT": contract.work_basis_text or "Акт технического осмотра",
            "WORK_BASIS_NUMBER": contract.work_basis_number or "",
            "WORK_BASIS_DATE": self._date_text(contract.work_basis_date) if contract.work_basis_date else "",
            "WORK_GOAL": contract.work_goal or "",
            "WORK_CONDITIONS": contract.work_conditions or "",
        }
        work_lines_data = []
        for index, line in enumerate(contract.work_lines.all().order_by("order", "id"), start=1):
            work_lines_data.append({
                "WORK_LINE_NO": str(index),
                "WORK_NAME": line.work_type or "",
                "WORK_UNIT": line.unit or "",
                "WORK_QUANTITY": str(line.quantity) if line.quantity else "",
            })
        if not work_lines_data:
            work_lines_data.append({
                "WORK_LINE_NO": "1",
                "WORK_NAME": contract.work_type or contract.subject,
                "WORK_UNIT": contract.volume_unit or "",
                "WORK_QUANTITY": str(contract.planned_volume) if contract.planned_volume else "",
            })
        context["__work_lines_data__"] = work_lines_data
        act_work_lines_data = []
        for item in work_lines_data:
            act_work_lines_data.append({
                "ACT_WORK_LINE_NO": item["WORK_LINE_NO"],
                "ACT_WORK_NAME": item["WORK_NAME"],
                "ACT_WORK_UNIT": item["WORK_UNIT"],
                "ACT_WORK_QUANTITY": item["WORK_QUANTITY"],
            })
        context["__act_work_lines_data__"] = act_work_lines_data
        return context

    def _supply_contract_template_context(self, contract: SupplyContract) -> dict[str, Any]:
        vat_rate = Decimal("20")
        vat_amount = Decimal(contract.amount or 0) * vat_rate / Decimal("100")
        profile = self._organization_profile()
        buyer_name = profile["name"] or "АО «СТ-1»"
        buyer_signer = profile.get("contractor_signer_name") or ""
        buyer_signer_position = profile.get("contractor_signer_position") or "Заместитель генерального директора"
        supplier_contact = contract.supplier.contact_person or ""
        return {
            **self._template_common_context(),
            **self._date_parts("CONTRACT", contract.contract_date),
            "CONTRACT_NUMBER": contract.number,
            "SUPPLIER_NAME": contract.supplier.name,
            "SUPPLIER_REQUISITES": self._supplier_requisites(contract.supplier) or "-",
            "SUPPLIER_SIGNER": supplier_contact or "________________",
            "SUPPLIER_AUTH_DOC": "доверенности",
            "SUPPLIER_SIGNER_POSITION": "Генеральный директор",
            "SUPPLIER_NAME_SHORT": contract.supplier.name,
            "BUYER_NAME": buyer_name,
            "BUYER_REQUISITES": self._organization_requisites() or "-",
            "BUYER_SIGNER": buyer_signer or "________________",
            "BUYER_AUTH_DOC": profile.get("contractor_auth_doc") or "доверенности",
            "BUYER_SIGNER_POSITION": buyer_signer_position,
            "BUYER_NAME_SHORT": buyer_name,
            "CONTRACT_AMOUNT": money(contract.amount),
            "VAT_RATE": str(vat_rate),
            "VAT_AMOUNT": money(vat_amount),
            "DELIVERY_PLACE": settings.WAREHOUSE_NAME,
            "DELIVERY_TERM": contract.terms or "по согласованным заявкам",
            "PAYMENT_TERM": "по условиям договора",
            "VALID_UNTIL": "",
            "LEFT_SIGNER_NAME": self._short_name(supplier_contact) if supplier_contact else contract.supplier.name,
            "RIGHT_SIGNER_NAME": self._short_name(buyer_signer) if buyer_signer else buyer_name,
        }

    def _primary_document_template_context(self, item: PrimaryDocument) -> dict[str, Any]:
        lines = list(item.lines.all())
        total_amount = Decimal(item.amount or 0)
        if not total_amount:
            total_amount = sum((self._line_amount(line) for line in lines), Decimal("0"))
        vat_amount = Decimal(item.vat_amount or 0)
        supplier_requisites = self._supplier_requisites(item.supplier) or "-"
        supplier_requisites_raw = "" if supplier_requisites == "-" else supplier_requisites
        buyer_name = self._organization_name() or "-"
        organization_profile = self._organization_profile()
        buyer_requisites = self._organization_requisites() or "-"
        supplier_inn = item.supplier.tax_id or ""
        supplier_kpp = self._extract_requisite(supplier_requisites_raw, r"КПП\s*[:№#-]?\s*([0-9]{9})")
        supplier_account = self._extract_requisite(supplier_requisites_raw, r"(?:р/с|расчетный\s+счет|расч[её]тный\s+счет)\s*[:№#-]?\s*([0-9]{20})")
        supplier_corr_account = self._extract_requisite(supplier_requisites_raw, r"(?:к/с|корр\.?\s*счет|корреспондентский\s+счет)\s*[:№#-]?\s*([0-9]{20})")
        supplier_bik = self._extract_requisite(supplier_requisites_raw, r"БИК\s*[:№#-]?\s*([0-9]{9})")
        organization_requisites_raw = self._organization_requisites()
        payer_account = self._extract_requisite(organization_requisites_raw, r"(?:р/с|расчетный\s+счет|расч[её]тный\s+счет)\s*[:№#-]?\s*([0-9]{20})")
        payer_corr_account = self._extract_requisite(organization_requisites_raw, r"(?:к/с|корр\.?\s*счет|корреспондентский\s+счет)\s*[:№#-]?\s*([0-9]{20})")
        payer_bik = self._extract_requisite(organization_requisites_raw, r"БИК\s*[:№#-]?\s*([0-9]{9})")
        if item.supply_contract and item.supply_contract.contract_date:
            contract_date = item.supply_contract.contract_date
            basis_document = (
                f'Договор № {item.supply_contract.number} от '
                f'"{contract_date.day:02d}" {MONTH_NAMES[contract_date.month - 1]} {contract_date.year} г.'
            )
        elif item.basis_reference:
            basis_document = item.basis_reference
        elif item.procurement_request:
            basis_document = item.procurement_request.number
        else:
            basis_document = ""
        context: dict[str, Any] = {
            **self._template_common_context(),
            "DOCUMENT_NUMBER": item.number,
            "DOCUMENT_DATE": self._date_text(item.doc_date),
            "INVOICE_NUMBER": item.number,
            "INVOICE_DATE": self._date_text(item.doc_date),
            "INVOICE_FACTURE_NUMBER": item.number,
            "INVOICE_FACTURE_DATE": self._date_text(item.doc_date),
            "WAYBILL_NUMBER": item.number,
            "WAYBILL_DATE": self._date_text(item.doc_date),
            "PAYMENT_ORDER_NUMBER": item.number,
            "PAYMENT_DATE": self._date_text(item.doc_date),
            "DEBITED_DATE": "",
            "RECEIVED_BY_BANK_DATE": "",
            "BASIS_DOCUMENT": basis_document,
            #"BASIS_DOCUMENT": item.basis_reference or (item.procurement_request.number if item.procurement_request else ""),
            "PAYMENT_DOCUMENT": item.basis_reference or "",
            "PAYMENT_PURPOSE": item.notes or item.basis_reference or f"Оплата по счету № {item.number}",
            "PAYMENT_PRIORITY": "5",
            "PAYMENT_KIND": "электронно",
            "PAYMENT_CODE": "",
            "OPERATION_TYPE": "01",
            "RESERVED_FIELD": "",
            "SUPPLIER_NAME": item.supplier.name,
            "SELLER_NAME": item.supplier.name,
            "SHIPPER_NAME": item.supplier.name,
            "SHIPPER_NAME_ADDRESS": f"{item.supplier.name}, {item.supplier.address or supplier_requisites_raw}".strip(", "),
            "SUPPLIER_REQUISITES": supplier_requisites,
            "SELLER_ADDRESS": item.supplier.address or supplier_requisites,
            "SUPPLIER_ADDRESS": item.supplier.address or supplier_requisites_raw,
            "SUPPLIER_PHONE": item.supplier.phone,
            "SHIPPER_REQUISITES": supplier_requisites,
            "SUPPLIER_INN": supplier_inn,
            "SUPPLIER_KPP": supplier_kpp,
            "SELLER_INN_KPP": supplier_inn,
            "BUYER_NAME": buyer_name,
            "BUYER_REQUISITES": buyer_requisites,
            "BUYER_ADDRESS": organization_profile.get("address", ""),
            "BUYER_INN": organization_profile.get("tax_id", ""),
            "BUYER_KPP": organization_profile.get("kpp", ""),
            "BUYER_INN_KPP": organization_profile.get("tax_id", ""),
            "CONSIGNEE_NAME": buyer_name,
            "CONSIGNEE_NAME_ADDRESS": f"{buyer_name}, {self._organization_profile().get('address', '')}".strip(", "),
            "CONSIGNEE_REQUISITES": buyer_requisites,
            "PAYER_NAME": buyer_name,
            "PAYER_INN": organization_profile.get("tax_id", ""),
            "PAYER_KPP": organization_profile.get("kpp", ""),
            "PAYER_ACCOUNT": payer_account,
            "PAYER_BANK_NAME": organization_profile.get("bank_details", ""),
            "PAYER_BANK_BIK": payer_bik,
            "PAYER_BANK_CORR_ACCOUNT": payer_corr_account,
            "PAYER_SIGNER_POSITION": "Директор",
            "PAYER_REQUISITES": buyer_requisites,
            "PAYEE_NAME": item.supplier.name,
            "PAYEE_INN": supplier_inn,
            "PAYEE_KPP": supplier_kpp,
            "PAYEE_ACCOUNT": supplier_account,
            "PAYEE_BANK_NAME": supplier_requisites_raw,
            "PAYEE_BANK_BIK": supplier_bik,
            "PAYEE_BANK_CORR_ACCOUNT": supplier_corr_account,
            "ITEMS_COUNT": len(lines),
            "AMOUNT": money(total_amount),
            "AMOUNT_WORDS": f"{money(total_amount)} руб.",
            "TOTAL_AMOUNT": money(total_amount),
            "TOTAL_TO_PAY": money(total_amount),
            "TOTAL_TO_PAY_WORDS": f"{money(total_amount)} руб.",
            "VAT_RATE": "20",
            "VAT_AMOUNT": money(vat_amount),
            "AMOUNT_NO_VAT": money(total_amount - vat_amount),
            "AMOUNT_WITH_VAT": money(total_amount),
            "CURRENCY_NAME": "Российский рубль",
            "CURRENCY_CODE": "643",
            "COUNTRY": "Россия",
            "EXCISE": "без акциза",
            "CUSTOMS_DECLARATION": "-",
            "PACKING_TYPE": "-",
            "PACKS_COUNT": "",
            "WEIGHT": "",
            "COMMENT": item.notes or "",
            "RECEIVER_BANK": supplier_requisites_raw,
            "BANK_BIK": supplier_bik,
            "BANK_ACCOUNT": supplier_account,
            "BANK_CORR_ACCOUNT": supplier_corr_account,
            "DIRECTOR_ORDER": "",
            "ACCOUNTANT_NAME": "",
        }
        lines_data = []
        for index, line in enumerate(lines, start=1):
            line_amount = self._line_amount(line)
            lines_data.append({
                "LINE_NO": str(index),
                "MATERIAL_CODE": line.material.code,
                "ITEM_CODE": line.material.code,
                "ITEM_NAME": line.material.name,
                "UNIT": line.material.unit,
                "QUANTITY": str(line.quantity),
                "PRICE": money(line.unit_price),
                "LINE_AMOUNT": money(line_amount),
                "AMOUNT_NO_VAT": money(line_amount),
                "AMOUNT_WITH_VAT": money(line_amount),
                "VAT_AMOUNT": "",
                "EXCISE": "без акциза",
                "VAT_RATE": "20%",
                "COUNTRY": "Россия",
                "CUSTOMS_DECLARATION": "-",
                "PACKING_TYPE": "-",
                "PACKS_COUNT": "",
                "WEIGHT": "",
            })
        context["__lines_data__"] = lines_data
        def map_line(line, index: int) -> dict[str, Any]:
            line_amount = self._line_amount(line)
            return {
                "LINE_NO": index,
                "MATERIAL_CODE": line.material.code,
                "ITEM_CODE": line.material.code,
                "ITEM_NAME": line.material.name,
                "UNIT": line.material.unit,
                "QUANTITY": line.quantity,
                "PRICE": money(line.unit_price),
                "LINE_AMOUNT": money(line_amount),
                "AMOUNT_NO_VAT": money(line_amount),
                "AMOUNT_WITH_VAT": money(line_amount),
                "VAT_AMOUNT": "",
            }
        self._add_line_context(context, lines, map_line)
        return context

    def _stock_receipt_template_context(self, receipt: StockReceipt) -> dict[str, Any]:
        profile = self._organization_profile()
        lines = list(receipt.lines.all())
        total_amount = sum((self._line_amount(line) for line in lines), Decimal("0"))
        total_qty = sum((Decimal(line.quantity or 0) for line in lines), Decimal("0"))
        supplier_document = receipt.supplier_document
        primary_document = receipt.primary_document
        context: dict[str, Any] = {
            **self._template_common_context(),
            "RECEIPT_ORDER_NUMBER": receipt.number,
            "DOCUMENT_DATE": self._date_text(receipt.receipt_date),
            "ORGANIZATION_NAME": self._organization_name() or "АО «СТ-1»",
            "STRUCTURAL_UNIT": settings.WAREHOUSE_NAME,
            "WAREHOUSE": settings.WAREHOUSE_NAME,
            "SUPPLIER_NAME": receipt.supplier.name,
            "SUPPLIER_DOCUMENT_NUMBER": supplier_document.doc_number if supplier_document else (primary_document.number if primary_document else ""),
            "SUPPLIER_DOCUMENT_DATE": self._date_text(supplier_document.doc_date if supplier_document else (primary_document.doc_date if primary_document else "")),
            "PAYMENT_DOCUMENT_NUMBER": primary_document.number if primary_document else "",
            "TOTAL_DOCUMENT_QTY": total_qty,
            "TOTAL_ACCEPTED_QTY": total_qty,
            "TOTAL_AMOUNT_NO_VAT": money(total_amount),
            "TOTAL_VAT": "",
            "TOTAL_WITH_VAT": money(total_amount),
            "OKPO": profile.get("okpo") or "",
            "ACCOUNT_CODE": "",
            "STOCK_CARD_NUMBER": "",
            "SENDER_POSITION": "Поставщик",
            "RECEIVER_POSITION": "Кладовщик",
            "LEFT_SIGNER_NAME": self._last_name_initials(receipt.supplier.contact_person) if receipt.supplier.contact_person else receipt.supplier.name,
            "RIGHT_SIGNER_NAME": self._last_name_initials(receipt.created_by.full_name_or_username) if receipt.created_by_id else "",
        }
        def map_line(line, index: int) -> dict[str, Any]:
            return {
                "MATERIAL_CODE": line.material.code,
                "MATERIAL_NAME": line.material.name,
                "UNIT": line.material.unit,
                "DOCUMENT_QTY": line.quantity,
                "ACCEPTED_QTY": line.quantity,
                "PRICE": money(line.unit_price),
                "AMOUNT_NO_VAT": money(self._line_amount(line)),
                "VAT_AMOUNT": "",
            }
        self._add_line_context(context, lines, map_line)
        return context

    def _stock_issue_template_context(self, issue: StockIssue) -> dict[str, Any]:
        from .models import User, RoleChoices
        lines = list(issue.lines.all())
        director = User.objects.filter(role=RoleChoices.DIRECTOR).first()
        context: dict[str, Any] = {
            **self._template_common_context(),
            "DOCUMENT_NUMBER": issue.number,
            "DOCUMENT_DATE": self._date_text(issue.issue_date),
            "SENDER_UNIT": settings.WAREHOUSE_NAME,
            "RECEIVER_UNIT": issue.site_name,
            "SENDER_ACTIVITY": "склад",
            "RECEIVER_ACTIVITY": "строительно-монтажные работы",
            "RECEIVED_BY_NAME": self._last_name_initials(issue.received_by_name) if issue.received_by_name else "",
            "RECEIVED_BY_POSITION": f"Начальник участка {issue.site_name}",
            "ISSUED_BY_POSITION": "Кладовщик",
            "OPERATION_CODE": "",
            "ACCOUNTING_UNIT": "",
            "ACCOUNT_CODE": "",
            "STOCK_RECEIPT_NUMBER": issue.stock_receipt.number if issue.stock_receipt else "",
            "STOCK_RECEIPT_DATE": self._date_text(issue.stock_receipt.receipt_date) if issue.stock_receipt else "",
            "THROUGH_WHOM": self._last_name_initials(issue.issued_by.full_name_or_username) if issue.issued_by_id else "",
            "REQUESTED_BY_NAME": self._last_name_initials(issue.received_by_name) if issue.received_by_name else "",
            "APPROVED_BY_NAME": self._last_name_initials(director.full_name_or_username) if director else "",
            "LEFT_SIGNER_NAME": self._last_name_initials(issue.issued_by.full_name_or_username) if issue.issued_by_id else "",
            "RIGHT_SIGNER_NAME": self._last_name_initials(issue.received_by_name) if issue.received_by_name else "",
        }
        def map_line(line, index: int) -> dict[str, Any]:
            return {
                "MATERIAL_CODE": line.material.code,
                "MATERIAL_NAME": line.material.name,
                "UNIT": line.material.unit,
                "REQUESTED_QTY": line.quantity,
                "ISSUED_QTY": line.quantity,
                "PRICE": money(line.unit_price),
                "AMOUNT": money(self._line_amount(line)),
                "STOCK_CARD_NUMBER": "",
            }
        self._add_line_context(context, lines, map_line)
        return context

    def _writeoff_template_context(self, act: WriteOffAct) -> dict[str, Any]:
        from .models import User, RoleChoices
        director = User.objects.filter(role=RoleChoices.DIRECTOR).first()
        site_manager = User.objects.filter(role=RoleChoices.SITE_MANAGER, site_name__iexact=act.site_name).first()
        context: dict[str, Any] = {
            **self._template_common_context(),
            **self._date_parts("APPROVAL", act.act_date),
            "ACT_NUMBER": act.number,
            "MONTH": MONTH_NAMES[act.act_date.month - 1],
            "YEAR": act.act_date.year,
            "SITE_NUMBER": act.site_name,
            "CONTRACT_NUMBER": act.contract.number if act.contract else "",
            "CONTRACT_DATE": self._date_text(act.contract.contract_date) if act.contract else "",
            "CONTRACT_SUBJECT": act.contract.subject if act.contract else "",
            "CONTRACT_WORK_COLUMN": act.work_type,
            "DIRECTOR_NAME": self._short_name(director.full_name_or_username) if director else "________________",
            # Комиссия: Фамилия И.О. (начальник участка — и как нач.участка, и как МОЛ)
            "SITE_MANAGER_NAME": self._last_name_initials(site_manager.full_name_or_username) if site_manager else "________________",
            "RESPONSIBLE_PERSON_NAME": self._last_name_initials(site_manager.full_name_or_username) if site_manager else "________________",
            # Подписи внизу: И.О.Фамилия
            "SITE_MANAGER_NAME_SHORT": self._short_name(site_manager.full_name_or_username) if site_manager else "________________",
            "RESPONSIBLE_PERSON_NAME_SHORT": self._short_name(site_manager.full_name_or_username) if site_manager else "________________",
        }
        return context

    def _ppe_template_context(self, issuance: PPEIssuance) -> dict[str, Any]:
        issued_by_name = issuance.issued_by.full_name_or_username if issuance.issued_by_id else ""
        confirmed_by_name = issuance.confirmed_by.full_name_or_username if issuance.confirmed_by_id else ""
        context: dict[str, Any] = {
            **self._template_common_context(),
            "DOCUMENT_NUMBER": issuance.number,
            "DOCUMENT_DATE": self._date_text(issuance.issue_date),
            "SITE_NAME": issuance.site_name,
            "PPE_CATEGORY": issuance.season or "СИЗ",
            "SITE_MANAGER_NAME": self._short_name(issued_by_name) if issued_by_name else "________________",
            "RESPONSIBLE_PERSON_NAME": self._short_name(confirmed_by_name) if confirmed_by_name else self._short_name(self._warehouse_user_name()),
            "LEFT_SIGNER_NAME": self._short_name(confirmed_by_name) if confirmed_by_name else self._short_name(self._warehouse_user_name()),
            "RIGHT_SIGNER_NAME": self._short_name(issued_by_name) if issued_by_name else "________________",
        }
        return context

    def _work_acceptance_template_context(self, act: WorkAcceptanceAct) -> dict[str, Any]:
        vat_rate = Decimal(act.contract.vat_rate or 0)
        vat_amount = Decimal(act.amount or 0) * vat_rate / Decimal("100")
        profile = self._organization_profile()
        customer_name = act.contract.resolved_customer_name() or "-"
        contractor_name = act.contract.resolved_contractor_name() or self._organization_name() or "-"
        customer_signer = act.contract.customer_signer_name or ""
        contractor_signer = profile.get("contractor_signer_name") or ""
        customer_signer_short = self._short_name(customer_signer) if customer_signer else "________________"
        contractor_signer_short = self._short_name(contractor_signer) if contractor_signer else "________________"
        return {
            **self._template_common_context(),
            **self._date_parts("ACT", act.act_date),
            **self._date_parts("CONTRACT", act.contract.contract_date),
            "ACT_NUMBER": act.number,
            "CONTRACT_NUMBER": act.contract.number,
            "CUSTOMER_NAME": customer_name,
            "CONTRACTOR_NAME": contractor_name,
            "OBJECT_ADDRESS": act.contract.object.address if act.contract.object else act.site_name,
            "WORK_DESCRIPTION": act.work_description or act.contract.subject,
            "WORK_PERIOD": f"{self._date_text(act.contract.start_date)} - {self._date_text(act.contract.end_date)}",
            "AMOUNT": money(act.amount),
            "VAT_RATE": vat_rate,
            "VAT_AMOUNT": money(vat_amount),
            "COPIES_COUNT": "2",
            "LEFT_SIGNER_NAME": customer_signer_short,
            "RIGHT_SIGNER_NAME": contractor_signer_short,
            "CUSTOMER_SIGNER_NAME": customer_signer_short,
            "CUSTOMER_SIGNER_NAME_FULL": customer_signer,
            "CUSTOMER_SIGNER_POSITION": act.contract.customer_signer_position or "представителя",
            "CUSTOMER_AUTH_DOC": act.contract.customer_auth_doc or "доверенности",
            "CONTRACTOR_SIGNER_NAME": contractor_signer_short,
            "CONTRACTOR_SIGNER_NAME_FULL": contractor_signer,
            "CONTRACTOR_SIGNER_POSITION": profile.get("contractor_signer_position") or "представителя",
            "CONTRACTOR_AUTH_DOC": profile.get("contractor_auth_doc") or "доверенности",
            "CUSTOMER_SIGNER_POSITION_GENITIVE": act.contract.customer_signer_position_genitive or act.contract.customer_signer_position or "представителя",
            "CUSTOMER_SIGNER_NAME_GENITIVE": act.contract.customer_signer_name_genitive or customer_signer or "________________",
            "CONTRACTOR_SIGNER_POSITION_GENITIVE": profile.get("contractor_signer_position_genitive") or profile.get("contractor_signer_position") or "представителя",
            "CONTRACTOR_SIGNER_NAME_GENITIVE": profile.get("contractor_signer_name_genitive") or contractor_signer or "________________",
        }

    def _supplier_document_template_name(self, doc_type: str) -> str | None:
        normalized = doc_type.casefold()
        if "счет-фактура" in normalized or "счёт-фактура" in normalized:
            return "Счет-фактура_шаблон.docx"
        if "наклад" in normalized:
            return "Товарная накладная ТОРГ-12_шаблон.docx"
        if "счет" in normalized or "счёт" in normalized:
            return "Счет на оплату по скану_шаблон.docx"
        return None

    def _export_smr_contract(self, entity_id: int) -> Path:
        contract = SMRContract.objects.select_related("object").get(pk=entity_id)
        path = self._doc_path("smr_contract", contract.number)
        context = self._smr_contract_template_context(contract)
        work_lines_data = context.pop("__work_lines_data__", [])
        act_work_lines_data = context.pop("__act_work_lines_data__", [])
        context["ACT_WORK_LINE_NO"] = "{{ACT_WORK_LINE_NO}}"
        context["ACT_WORK_NAME"] = "{{ACT_WORK_NAME}}"
        context["ACT_WORK_UNIT"] = "{{ACT_WORK_UNIT}}"
        context["ACT_WORK_QUANTITY"] = "{{ACT_WORK_QUANTITY}}"
        tmp_path = path.with_suffix(".tmp.docx")
        ok = self._render_docx_template_with_table_rows(
            DOCX_TEMPLATE_FILES["smr_contract"], context, tmp_path,
            table_rows=work_lines_data, row_marker="{{WORK_LINE_NO}}",
        )
        if ok:
            self._render_docx_file_with_table_rows(
                tmp_path, {}, path,
                table_rows=act_work_lines_data, row_marker="{{ACT_WORK_LINE_NO}}",
            )
            tmp_path.unlink(missing_ok=True)
            return path
        customer_name = contract.resolved_customer_name() or "-"
        customer_requisites = contract.resolved_customer_requisites() or "-"
        contractor_name = contract.resolved_contractor_name() or "-"
        contractor_requisites = contract.resolved_contractor_requisites() or "-"
        doc = self._prepare_doc("ДОГОВОР НА ВЫПОЛНЕНИЕ СМР", f"№ {contract.number} от {contract.contract_date}")
        self._add_meta(doc, [
            ("Заказчик", customer_name),
            ("Реквизиты заказчика", customer_requisites),
            ("Подрядчик", contractor_name),
            ("Реквизиты подрядчика", contractor_requisites),
            ("Объект", contract.object.name if contract.object else ""),
            ("Предмет", contract.subject),
            ("Вид работ", contract.work_type or ""),
            ("Плановый объем", f"{contract.planned_volume or 0} {contract.volume_unit or ''}".strip()),
            ("Стоимость", f"{money(contract.amount)} руб."),
            ("Сроки", f"{contract.start_date or '-'} - {contract.end_date or '-'}"),
        ])
        self._add_heading(doc, "1. Предмет договора")
        self._add_clause(doc, "1.1", f"Подрядчик обязуется выполнить работы «{contract.subject}» на объекте «{contract.object.name if contract.object else '-'}», а Заказчик обязуется принять и оплатить результат работ.")
        self._add_clause(doc, "1.2", f"Вид работ: {contract.work_type or '-'}. Плановый объем: {contract.planned_volume or '-'} {contract.volume_unit or ''}.")
        self._add_heading(doc, "2. Стоимость и сроки")
        self._add_clause(doc, "2.1", f"Стоимость работ составляет {money(contract.amount)} руб., НДС {contract.vat_rate}%.")
        self._add_clause(doc, "2.2", f"Срок выполнения работ: с {contract.start_date or '-'} по {contract.end_date or '-'}.")
        self._add_signature(doc, f"Заказчик: {customer_name}", f"Подрядчик: {contractor_name}")
        doc.save(path)
        return path

    def _export_supply_contract(self, entity_id: int) -> Path:
        contract = SupplyContract.objects.select_related("supplier", "related_smr_contract").get(pk=entity_id)
        path = self._doc_path("supply_contract", contract.number)
        if self._render_docx_template(DOCX_TEMPLATE_FILES["supply_contract"], self._supply_contract_template_context(contract), path):
            return path
        buyer_name = self._organization_name() or "-"
        buyer_requisites = self._organization_requisites() or "-"
        supplier_requisites = self._supplier_requisites(contract.supplier) or "-"
        doc = self._prepare_doc("ДОГОВОР ПОСТАВКИ", f"№ {contract.number} от {contract.contract_date}")
        self._add_meta(doc, [
            ("Поставщик", contract.supplier.name),
            ("Реквизиты поставщика", supplier_requisites),
            ("Покупатель", buyer_name),
            ("Реквизиты покупателя", buyer_requisites),
            ("Связанный договор СМР", contract.related_smr_contract.number if contract.related_smr_contract else "-"),
            ("Сумма", f"{money(contract.amount)} руб."),
            ("Статус", contract.get_status_display()),
        ])
        doc.add_paragraph(contract.terms or "Поставка материалов выполняется по заявкам снабженца через АИС.")
        self._add_signature(doc, f"Поставщик: {contract.supplier.name}", f"Покупатель: {buyer_name}")
        doc.save(path)
        return path

    def _export_site_material_request(self, entity_id: int) -> Path:
        from .models import User, RoleChoices
        request = SiteMaterialRequest.objects.select_related("contract", "requested_by").prefetch_related("lines__material").get(pk=entity_id)
        requester_name = request.requested_by.full_name_or_username if request.requested_by_id else ""
        requester_short = self._last_name_initials(requester_name) if requester_name else "________________"
        warehouse_user = User.objects.filter(role=RoleChoices.WAREHOUSE, is_active=True).first()
        warehouse_short = self._last_name_initials(warehouse_user.full_name_or_username) if warehouse_user else "________________"
        doc = self._prepare_doc("ЗАЯВКА НА МАТЕРИАЛЫ СО СКЛАДА", f"№ {request.number} от {request.request_date}")
        self._add_meta(doc, [
            ("Участок", request.site_name),
            ("Договор СМР", request.contract.number if request.contract else "-"),
            ("Заявитель", requester_name),
            ("Статус", request.get_status_display()),
        ])
        self._add_table(
            doc,
            ["Код", "Наименование", "Ед.", "По договору", "Запас", "Итого", "Примечание"],
            [
                [
                    line.material.code,
                    line.material.name,
                    line.material.unit,
                    str(line.quantity - line.reserve_qty),
                    str(line.reserve_qty),
                    str(line.quantity),
                    line.notes or "",
                ]
                for line in request.lines.all()
            ],
        )
        self._add_signature_borderless(
            doc,
            "Начальник участка",
            "Кладовщик",
            requester_short,
            warehouse_short,
        )
        path = self._doc_path("site_material_request", request.number)
        doc.save(path)
        return path

    def _export_procurement_request(self, entity_id: int) -> Path:
        request = ProcurementRequest.objects.select_related("contract", "site_request", "supplier", "requested_by").prefetch_related("lines__material").get(pk=entity_id)
        doc = self._prepare_doc("ЗАЯВКА НА ЗАКУПКУ МАТЕРИАЛОВ", f"№ {request.number} от {request.request_date}")
        self._add_meta(doc, [
            ("Участок", request.site_name),
            ("Договор СМР", request.contract.number if request.contract else "-"),
            ("Основание", f"Заявка участка {request.site_request.number}" if request.site_request else "-"),
            ("Поставщик", request.supplier.name if request.supplier else "-"),
            ("Статус", request.get_status_display()),
        ])
        self._add_table(
            doc,
            ["Код", "Наименование", "Ед.", "Количество", "Примечание"],
            [
                [
                    line.material.code,
                    line.material.name,
                    line.material.unit,
                    str(line.quantity),
                    line.notes or "",
                ]
                for line in request.lines.all()
            ],
        )
        supplier_name = (
            self._last_name_initials(request.supplier.contact_person)
            if request.supplier and request.supplier.contact_person
            else (request.supplier.name if request.supplier else "________________")
        )
        requester_name = (
            self._last_name_initials(request.requested_by.full_name_or_username)
            if request.requested_by_id
            else "________________"
        )
        self._add_signature_borderless(
            doc,
            "Снабженец",
            "Поставщик",
            requester_name,
            supplier_name,
        )
        path = self._doc_path("procurement_request", request.number)
        doc.save(path)
        return path

    def _export_primary_document(self, entity_id: int) -> Path:
        item = (
            PrimaryDocument.objects.select_related("document_type", "supplier", "procurement_request", "supply_contract", "stock_receipt")
            .prefetch_related("lines__material")
            .get(pk=entity_id)
        )
        context = self._primary_document_template_context(item)
        lines_data = context.pop("__lines_data__", [])
        xlsx_template_name = PRIMARY_DOCUMENT_XLSX_TEMPLATE_FILES.get(item.document_type.code)
        if xlsx_template_name:
            path = self._xlsx_document_path(item.document_type.code, item.number)
            if self._render_xlsx_template(xlsx_template_name, context, path):
                return path
        path = self._doc_path(item.document_type.code, item.number)
        template_name = PRIMARY_DOCUMENT_TEMPLATE_FILES.get(item.document_type.code)
        if template_name:
            if lines_data:
                if self._render_docx_template_with_table_rows(template_name, context, path, table_rows=lines_data, row_marker="{{LINE_NO}}"):
                    return path
            else:
                if self._render_docx_template(template_name, context, path):
                    return path
        receiver_name = self._organization_name() or "-"
        receiver_requisites = self._organization_requisites() or "-"
        supplier_requisites = self._supplier_requisites(item.supplier) or "-"
        title_map = {
            "invoice": "СЧЕТ НА ОПЛАТУ",
            "invoice_facture": "СЧЕТ-ФАКТУРА",
            "upd": "УПД",
            "vat_invoice": "СЧЕТ-ФАКТУРА",
            "goods_waybill": "ТОВАРНАЯ НАКЛАДНАЯ",
            "receipt_invoice": "ПРИХОДНАЯ НАКЛАДНАЯ",
            "payment_order": "ПЛАТЕЖНОЕ ПОРУЧЕНИЕ",
        }
        doc = self._prepare_doc(title_map.get(item.document_type.code, item.document_type.name.upper()), f"№ {item.number} от {item.doc_date}")
        self._add_meta(doc, [
            ("Тип документа", item.document_type.name),
            ("Поставщик", item.supplier.name),
            ("Реквизиты поставщика", supplier_requisites),
            ("Получатель", receiver_name),
            ("Реквизиты получателя", receiver_requisites),
            ("Основание", item.basis_reference or "-"),
            ("Участок/склад", item.site_name or "-"),
            ("Сумма", f"{money(item.amount)} руб."),
            ("НДС", f"{money(item.vat_amount)} руб."),
            ("Статус", item.get_status_display()),
        ])
        self._add_table(
            doc,
            ["Код", "Наименование", "Ед.", "Количество", "Цена", "Сумма", "Примечание"],
            [
                [
                    line.material.code, line.material.name, line.material.unit,
                    str(line.quantity), money(line.unit_price),
                    money(line.quantity * line.unit_price), line.notes or "",
                ]
                for line in item.lines.all()
            ],
        )
        self._add_signature(doc, f"Поставщик: {item.supplier.name}", f"Получатель: {receiver_name}")
        doc.save(path)
        return path

    def _export_stock_receipt(self, entity_id: int) -> Path:
        receipt = StockReceipt.objects.select_related("supplier", "supplier_document", "primary_document").prefetch_related("lines__material").get(pk=entity_id)
        path = self._doc_path("stock_receipt", receipt.number)
        context = self._stock_receipt_template_context(receipt)
        lines_data = []
        total_no_vat = Decimal("0")
        total_vat = Decimal("0")
        total_with_vat = Decimal("0")
        for line in receipt.lines.all():
            line_amount = self._line_amount(line)
            line_vat = line_amount * Decimal("20") / Decimal("100")
            line_with_vat = line_amount + line_vat
            total_no_vat += line_amount
            total_vat += line_vat
            total_with_vat += line_with_vat
            lines_data.append({
                "MATERIAL_CODE": line.material.code,
                "MATERIAL_NAME": line.material.name,
                "UNIT": line.material.unit,
                "DOCUMENT_QTY": str(line.quantity),
                "ACCEPTED_QTY": str(line.quantity),
                "PRICE": money(line.unit_price),
                "AMOUNT_NO_VAT": money(line_amount),
                "VAT_AMOUNT": money(line_vat),
                "AMOUNT_WITH_VAT": money(line_with_vat),
            })
        context["TOTAL_AMOUNT_NO_VAT"] = money(total_no_vat)
        context["TOTAL_VAT"] = money(total_vat)
        context["TOTAL_WITH_VAT"] = money(total_with_vat)
        if lines_data and self._render_docx_template_with_table_rows(
            DOCX_TEMPLATE_FILES["stock_receipt"], context, path,
            table_rows=lines_data, row_marker="{{MATERIAL_CODE}}",
        ):
            return path
        if self._render_docx_template(DOCX_TEMPLATE_FILES["stock_receipt"], context, path):
            return path
        doc = self._prepare_doc("ПРИХОДНЫЙ ОРДЕР", f"№ {receipt.number} от {receipt.receipt_date}")
        self._add_meta(doc, [
            ("Поставщик", receipt.supplier.name),
            ("Документ поставщика", receipt.supplier_document.doc_number if receipt.supplier_document else "-"),
            ("Товарная накладная / УПД", receipt.primary_document.number if receipt.primary_document else "-"),
            ("Склад", settings.WAREHOUSE_NAME),
            ("Статус", receipt.get_status_display()),
        ])
        self._add_table(
            doc,
            ["Код", "Наименование", "Ед.", "Количество", "Цена", "Сумма"],
            [
                [
                    line.material.code, line.material.name, line.material.unit,
                    str(line.quantity), money(line.unit_price),
                    money(line.quantity * line.unit_price),
                ]
                for line in receipt.lines.all()
            ],
        )
        self._add_signature(doc, "Кладовщик", "Материально ответственное лицо")
        doc.save(path)
        return path

    def _export_stock_issue(self, entity_id: int) -> Path:
        issue = StockIssue.objects.select_related("contract", "site_request", "stock_receipt").prefetch_related("lines__material").get(pk=entity_id)
        path = self._doc_path("stock_issue", issue.number)
        context = self._stock_issue_template_context(issue)
        lines_data = []
        for line in issue.lines.all():
            lines_data.append({
                "MATERIAL_NAME": line.material.name,
                "MATERIAL_CODE": line.material.code,
                "UNIT": line.material.unit,
                "REQUESTED_QTY": str(line.quantity),
                "ISSUED_QTY": str(line.quantity),
                "PRICE": money(line.unit_price),
                "AMOUNT": money(self._line_amount(line)),
                "STOCK_CARD_NUMBER": "",
            })
        if lines_data and self._render_docx_template_with_table_rows(
            DOCX_TEMPLATE_FILES["stock_issue"], context, path,
            table_rows=lines_data, row_marker="{{MATERIAL_NAME}}",
        ):
            return path
        if self._render_docx_template(DOCX_TEMPLATE_FILES["stock_issue"], context, path):
            return path
        doc = self._prepare_doc("ТРЕБОВАНИЕ-НАКЛАДНАЯ", f"№ {issue.number} от {issue.issue_date}")
        self._add_meta(doc, [
            ("Участок", issue.site_name),
            ("Договор СМР", issue.contract.number if issue.contract else "-"),
            ("Заявка участка", issue.site_request.number if issue.site_request else "-"),
            ("Получатель", issue.received_by_name),
            ("Статус", issue.get_status_display()),
        ])
        self._add_table(
            doc,
            ["Код", "Наименование", "Ед.", "Количество", "Цена"],
            [
                [
                    line.material.code, line.material.name, line.material.unit,
                    str(line.quantity), money(line.unit_price),
                ]
                for line in issue.lines.all()
            ],
        )
        self._add_signature(doc, "Кладовщик", "Начальник участка")
        doc.save(path)
        return path

    def _export_writeoff(self, entity_id: int) -> Path:
        act = WriteOffAct.objects.select_related("contract__object").prefetch_related("lines__material").get(pk=entity_id)
        path = self._doc_path("write_off", act.number)
        template_key = (
            "write_off_production_economic"
            if act.template_variant == WriteOffTemplateVariant.PRODUCTION_ECONOMIC
            else "write_off"
        )
        context = self._writeoff_template_context(act)
        lines_data = [
            {
                "LINE_NO": str(i),
                "MATERIAL_CODE": line.material.code,
                "MATERIAL_NAME": line.material.name,
                "REPORT_MATERIAL_NAME": line.material.name,
                "UNIT": line.material.unit,
                "WORK_OR_FORM": act.work_type,
                "ACTUAL_QTY": str(line.actual_quantity),
                "NORM_QTY": str(line.calculated_quantity),
                "QUANTITY": str(line.actual_quantity),
                "NOTE": line.notes or "",
            }
            for i, line in enumerate(act.lines.all(), start=1)
        ]
        if lines_data and self._render_docx_template_with_table_rows(
            DOCX_TEMPLATE_FILES[template_key], context, path,
            table_rows=lines_data, row_marker="{{LINE_NO}}",
        ):
            return path
        if self._render_docx_template(DOCX_TEMPLATE_FILES[template_key], context, path):
            return path
        doc = self._prepare_doc("АКТ СПИСАНИЯ МАТЕРИАЛОВ", f"№ {act.number} от {act.act_date}")
        self._add_meta(doc, [
            ("Договор", act.contract.number if act.contract else "-"),
            ("Объект", act.contract.object.name if act.contract and act.contract.object else ""),
            ("Участок", act.site_name),
            ("Вид работ", act.work_type),
            ("Объем работ", f"{act.work_volume} {act.volume_unit}".strip()),
            ("Статус", act.get_status_display()),
        ])
        self._add_table(
            doc,
            ["Код", "Наименование", "Ед.", "Норма", "Расчет", "Факт"],
            [
                [
                    line.material.code, line.material.name, line.material.unit,
                    str(line.norm_per_unit), str(line.calculated_quantity), str(line.actual_quantity),
                ]
                for line in act.lines.all()
            ],
        )
        self._add_signature(doc, "Начальник участка", "Начальник монтажного объекта")
        doc.save(path)
        return path

    def _export_work_acceptance(self, entity_id: int) -> Path:
        act = WorkAcceptanceAct.objects.select_related("contract__object").get(pk=entity_id)
        path = self._doc_path("work_acceptance", act.number)
        if self._render_docx_template(DOCX_TEMPLATE_FILES["work_acceptance"], self._work_acceptance_template_context(act), path):
            return path
        customer_name = act.contract.resolved_customer_name() or act.contract.customer_name or "-"
        contractor_name = act.contract.resolved_contractor_name() or self._organization_name() or "-"
        doc = self._prepare_doc("АКТ СДАЧИ-ПРИЕМКИ ВЫПОЛНЕННЫХ РАБОТ", f"№ {act.number} от {act.act_date}")
        self._add_meta(doc, [
            ("Договор СМР", act.contract.number),
            ("Объект", act.contract.object.name if act.contract.object else act.site_name),
            ("Заказчик", customer_name),
            ("Подрядчик", contractor_name),
            ("Описание работ", act.work_description or act.contract.subject),
            ("Принятый объем", f"{act.accepted_volume or '-'} {act.volume_unit or ''}".strip()),
            ("Сумма", f"{money(act.amount)} руб."),
            ("Статус", act.get_status_display()),
        ])
        self._add_clause(doc, "1", "Работы выполнены в соответствии с договором и переданы Заказчику для приемки.")
        self._add_clause(doc, "2", "Подписание акта подтверждает закрытие выполненного этапа работ по договору СМР.")
        self._add_signature(doc, f"Заказчик: {customer_name}", f"Подрядчик: {contractor_name}")
        doc.save(path)
        return path

    def _export_ppe_issuance(self, entity_id: int) -> Path:
        issuance = PPEIssuance.objects.select_related("issued_by", "confirmed_by").prefetch_related("lines__worker", "lines__material").get(pk=entity_id)
        path = self._doc_path("ppe_issuance", issuance.number)
        season = (issuance.season or "").casefold()
        if "зим" in season:
            template_name = "Ведомость выдачи спецодежды зимняя_шаблон.docx"
        elif "лет" in season:
            template_name = "Ведомость выдачи спецодежды летняя_шаблон.docx"
        else:
            template_name = "Ведомость выдачи спецодежды перчатки_шаблон.docx"
        context = self._ppe_template_context(issuance)
        lines = list(issuance.lines.select_related("worker", "material").all())
        worker_numbers: dict[int, int] = {}
        worker_counter = 0
        for line in lines:
            if line.worker_id not in worker_numbers:
                worker_counter += 1
                worker_numbers[line.worker_id] = worker_counter
        lines_data = []
        for line in lines:
            sizes = []
            if line.clothing_size:
                sizes.append(f"размер одежды {line.clothing_size}")
            if line.shoe_size:
                sizes.append(f"размер обуви {line.shoe_size}")
            size_note = f" ({', '.join(sizes)})" if sizes else ""
            lines_data.append({
                "LINE_NO": str(worker_numbers[line.worker_id]),
                "WORKER_NAME": line.worker.full_name,
                "EMPLOYEE_NUMBER": line.worker.employee_number,
                "PPE_NAME": f"{line.material.name}{size_note}",
                "PPE_CODE": line.material.code,
                "NOMENCLATURE_NUMBER": line.material.code,
                "UNIT": line.material.unit,
                "UNIT_NAME": line.material.unit,
                "UNIT_CODE": line.material.unit,
                "QUANTITY": str(int(line.quantity)),
                "SERVICE_LIFE_MONTHS": str(line.service_life_months),
                "START_DATE": self._date_text(line.replacement_start_date),
                "WORKER_SIGNATURE": "",
            })
        if lines_data and self._render_docx_template_with_table_rows(
            template_name, context, path,
            table_rows=lines_data, row_marker="{{LINE_NO}}",
        ):
            return path
        if self._render_docx_template(template_name, context, path):
            return path
        doc = self._prepare_doc("ВЕДОМОСТЬ ВЫДАЧИ СПЕЦОДЕЖДЫ", f"№ {issuance.number} от {issuance.issue_date}")
        self._add_meta(doc, [
            ("Участок", issuance.site_name),
            ("Сезон", issuance.season),
            ("Статус", issuance.get_status_display()),
        ])
        self._add_table(
            doc,
            ["Таб.№", "ФИО", "Материал", "Код", "Размер одежды", "Размер обуви", "Ед.", "Кол-во", "Срок службы, мес."],
            [
                [
                    line.worker.employee_number, line.worker.full_name,
                    line.material.name, line.material.code,
                    line.clothing_size, line.shoe_size,
                    line.material.unit, str(line.quantity), str(line.service_life_months),
                ]
                for line in issuance.lines.all()
            ],
        )
        self._add_signature(doc, "Материально ответственное лицо", "Начальник участка")
        doc.save(path)
        return path

    def _export_supplier_document(self, entity_id: int) -> Path:
        item = SupplierDocument.objects.select_related("supplier", "request", "supply_contract").get(pk=entity_id)
        path = self._doc_path("supplier_document", item.doc_number)
        template_name = self._supplier_document_template_name(item.doc_type)
        lines_data = []
        from .models import SupplierDocumentLine
        invoice_price_index: dict[str, Decimal] = {}
        normalized_doc_type = (item.doc_type or "").casefold()
        if "наклад" in normalized_doc_type or "упд" in normalized_doc_type or "фактур" in normalized_doc_type:
            invoice_qs = SupplierDocument.objects.filter(
                supplier=item.supplier,
                doc_type__in=["Счёт", "Счет", "Счёт на оплату", "Счет на оплату"],
            )
            if item.request_id:
                invoice_qs = invoice_qs.filter(request_id=item.request_id)
            elif item.supply_contract_id:
                invoice_qs = invoice_qs.filter(supply_contract_id=item.supply_contract_id)
            invoice = invoice_qs.order_by("-doc_date", "-id").first()
            if invoice:
                for sdl in SupplierDocumentLine.objects.select_related("material").filter(document=invoice):
                    if sdl.unit_price:
                        invoice_price_index[sdl.material.code] = Decimal(sdl.unit_price)
        supplier_lines = list(SupplierDocumentLine.objects.select_related("material").filter(document=item))
        if supplier_lines:
            for index, line in enumerate(supplier_lines, start=1):
                unit_price = Decimal(line.unit_price or 0)
                if unit_price == 0 and line.material.code in invoice_price_index:
                    unit_price = invoice_price_index[line.material.code]
                line_amount = Decimal(line.quantity or 0) * unit_price
                vat_rate = Decimal(item.vat_rate or 0)
                line_vat = line_amount * vat_rate / Decimal("100")
                lines_data.append({
                    "LINE_NO": str(index),
                    "ITEM_NAME": line.material.name,
                    "ITEM_CODE": line.material.code,
                    "MATERIAL_CODE": line.material.code,
                    "MATERIAL_NAME": line.material.name,
                    "UNIT": line.material.unit,
                    "PACKING_TYPE": "-",
                    "PACKS_COUNT": "",
                    "WEIGHT": "",
                    "QUANTITY": str(line.quantity),
                    "PRICE": money(unit_price),
                    "LINE_AMOUNT": money(line_amount),
                    "AMOUNT_NO_VAT": money(line_amount),
                    "VAT_AMOUNT": money(line_vat),
                    "AMOUNT_WITH_VAT": money(line_amount + line_vat),
                })
        elif item.request:
            for index, line in enumerate(item.request.lines.select_related("material").all(), start=1):
                req_unit_price = Decimal(line.unit_price or 0)
                line_amount = Decimal(line.quantity or 0) * req_unit_price
                lines_data.append({
                    "LINE_NO": str(index),
                    "ITEM_NAME": line.material.name,
                    "ITEM_CODE": line.material.code,
                    "MATERIAL_CODE": line.material.code,
                    "MATERIAL_NAME": line.material.name,
                    "UNIT": line.material.unit,
                    "PACKING_TYPE": "-",
                    "PACKS_COUNT": "",
                    "WEIGHT": "",
                    "QUANTITY": str(line.quantity),
                    "PRICE": money(req_unit_price),
                    "LINE_AMOUNT": money(line_amount),
                    "AMOUNT_NO_VAT": money(line_amount),
                    "VAT_AMOUNT": "",
                    "AMOUNT_WITH_VAT": money(line_amount),
                })
        if template_name:
            from .models import User, RoleChoices
            _warehouse_user = User.objects.filter(role=RoleChoices.WAREHOUSE, is_active=True).first()
            _warehouse_short = self._last_name_initials(_warehouse_user.full_name_or_username) if _warehouse_user else ""

            supplier_requisites = self._supplier_requisites(item.supplier) or "-"
            buyer_name = self._organization_name() or "-"
            buyer_requisites = self._organization_requisites() or "-"
            uploaded_by_name = item.uploaded_by.full_name_or_username if item.uploaded_by_id else ""
            org_profile = self._organization_profile()

            _supplier_signer = self._last_name_initials(item.supplier.contact_person) if item.supplier.contact_person else item.supplier.name
            _supplier_is_ip = len((item.supplier.tax_id or "").strip()) == 12
            if "наклад" in normalized_doc_type or "упд" in normalized_doc_type:
                _left_signer = _supplier_signer
                _right_signer = _warehouse_short
            else:
                _left_signer = _supplier_signer
                _right_signer = self._short_name(uploaded_by_name) if uploaded_by_name else _supplier_signer

            if item.supply_contract and item.supply_contract.contract_date:
                contract_date = item.supply_contract.contract_date
                basis_document = (
                    f'Договор № {item.supply_contract.number} от '
                    f'"{contract_date.day:02d}" {MONTH_NAMES[contract_date.month - 1]} {contract_date.year} г.'
                )
            elif item.request:
                basis_document = item.request.number
            else:
                basis_document = ""

            context = {
                **self._template_common_context(),
                "INVOICE_NUMBER": item.doc_number,
                "RECEIVER_BANK": (
                    self._extract_requisite(supplier_requisites, r"р/с\s+\S+\s+в\s+([^;,\n]+)")
                    or self._extract_requisite(supplier_requisites, r"(?:банк|в\s+банке?)\s*[:—]?\s*([^;,\n]+)")
                    or ""
                ),
                "BANK_BIK": self._extract_requisite(supplier_requisites, r"БИК\s*[:№#-]?\s*([0-9]{9})"),
                "BANK_ACCOUNT": self._extract_requisite(supplier_requisites, r"(?:р/с|расчетный\s+счет|расч[её]тный\s+счет)\s*[:№#-]?\s*([0-9]{20})"),
                "BANK_CORR_ACCOUNT": self._extract_requisite(supplier_requisites, r"(?:к/с|корр\.?\s*счет|корреспондентский\s+счет)\s*[:№#-]?\s*([0-9]{20})"),
                "SUPPLIER_INN": item.supplier.tax_id or "",
                "SUPPLIER_KPP": self._extract_requisite(supplier_requisites, r"КПП\s*[:№#-]?\s*([0-9]{9})"),
                "SUPPLIER_ADDRESS": item.supplier.address or "",
                "INVOICE_DATE": self._date_text(item.doc_date),
                "INVOICE_FACTURE_NUMBER": item.doc_number,
                "INVOICE_FACTURE_DATE": self._date_text(item.doc_date),
                "WAYBILL_NUMBER": item.doc_number,
                "WAYBILL_DATE": self._date_text(item.doc_date),
                "SUPPLIER_NAME": item.supplier.name,
                "SELLER_NAME": item.supplier.name,
                "SHIPPER_NAME": item.supplier.name,
                "SHIPPER_NAME_ADDRESS": f"{item.supplier.name}, {item.supplier.address or supplier_requisites}".strip(", "),
                "SUPPLIER_REQUISITES": supplier_requisites,
                "SHIPPER_REQUISITES": supplier_requisites,
                "BUYER_NAME": buyer_name,
                "BUYER_INN": org_profile.get("tax_id") or "",
                "BUYER_KPP": org_profile.get("kpp") or "",
                "BUYER_ADDRESS": org_profile.get("address") or "",
                "CONSIGNEE_NAME": buyer_name,
                "CONSIGNEE_NAME_ADDRESS": f"{buyer_name}, {self._organization_profile().get('address', '')}".strip(", "),
                "CONSIGNEE_REQUISITES": buyer_requisites,
                "PAYER_NAME": buyer_name,
                "PAYER_REQUISITES": buyer_requisites,
                "BUYER_REQUISITES": buyer_requisites,
                "TOTAL_AMOUNT": money(Decimal(item.amount or 0) - Decimal(item.vat_amount or 0)),
                "AMOUNT": money(Decimal(item.amount or 0) - Decimal(item.vat_amount or 0)),
                "AMOUNT_NO_VAT": money(Decimal(item.amount or 0) - Decimal(item.vat_amount or 0)),
                "AMOUNT_WITH_VAT": money(item.amount),
                "TOTAL_TO_PAY": money(item.amount),
                "TOTAL_TO_PAY_WORDS": f"{money(item.amount)} руб.",
                "VAT_AMOUNT": money(item.vat_amount),
                "VAT_RATE": f"{int(item.vat_rate or 20)}%",
                "BASIS_DOCUMENT": basis_document,
                "ITEMS_COUNT": len(lines_data),
                "TOTAL_AMOUNT_NO_VAT": money(Decimal(item.amount or 0) - Decimal(item.vat_amount or 0)),
                "TOTAL_AMOUNT_WITH_VAT": money(item.amount),
                "LEFT_SIGNER_NAME": _left_signer,
                "RIGHT_SIGNER_NAME": _right_signer,
                "ORG_HEAD_NAME": "" if _supplier_is_ip else _supplier_signer,
                "IP_NAME": item.supplier.name if _supplier_is_ip else "",
                "IP_OGRNIP": (item.supplier.ogrnip or "") if _supplier_is_ip else "",
                "SUPPLY_CONTRACT_NUMBER": item.supply_contract.number if item.supply_contract else "",
                "SUPPLY_CONTRACT_DATE": self._date_text(item.supply_contract.contract_date) if item.supply_contract else "",
                "PAYER_ACCOUNT": self._organization_profile().get("account", ""),
                "PAYER_BANK_NAME": self._organization_profile().get("bank_name", ""),
                "PAYER_BANK_BIK": self._organization_profile().get("bik", ""),
                "PAYER_BANK_CORR_ACCOUNT": self._organization_profile().get("corr_account", ""),
            }
            if lines_data:
                if self._render_docx_template_with_table_rows(
                    template_name, context, path,
                    table_rows=lines_data, row_marker="{{LINE_NO}}",
                ):
                    return path
            else:
                if self._render_docx_template(template_name, context, path):
                    return path
        doc = self._prepare_doc("ДОКУМЕНТ ПОСТАВКИ", f"{item.doc_type} № {item.doc_number} от {item.doc_date}")
        self._add_meta(doc, [
            ("Поставщик", item.supplier.name),
            ("Сумма", f"{money(item.amount)} руб."),
            ("НДС", f"{money(item.vat_amount)} руб."),
            ("Комментарий", item.notes or ""),
        ])
        doc.save(path)
        return path